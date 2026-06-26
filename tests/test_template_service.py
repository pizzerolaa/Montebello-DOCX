from __future__ import annotations

from zipfile import ZipFile

import pytest
from docx import Document

from app.exceptions import MarkerNotFoundError
from app.services.template_service import ParagraphSpec, insert_paragraphs_at_marker, replace_placeholders


def test_marker_insertion_replaces_marker_and_placeholders(workspace_tmp_path) -> None:
    path = workspace_tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Hola {{NAME}}")
    document.add_paragraph("[[MARKER]]")
    document.save(path)

    document = Document(path)
    replace_placeholders(document, {"NAME": "Mundo"})
    insert_paragraphs_at_marker(document, "[[MARKER]]", [ParagraphSpec("Insertado", bold=True)])

    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "Hola Mundo" in texts
    assert "Insertado" in texts
    assert "[[MARKER]]" not in texts
    inserted = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Insertado")
    assert inserted.paragraph_format.space_after.pt == 0
    assert inserted.paragraph_format.line_spacing == 1.0


def test_marker_insertion_finds_table_cell_marker() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "[[MARKER]]"
    insert_paragraphs_at_marker(document, "[[MARKER]]", [ParagraphSpec("Desde tabla")])
    assert "Desde tabla" in [paragraph.text for paragraph in table.cell(0, 0).paragraphs]


def test_marker_insertion_supports_inline_bold_highlight_and_numbering(workspace_tmp_path) -> None:
    path = workspace_tmp_path / "numbered.docx"
    document = Document()
    document.add_paragraph("[[MARKER]]")
    insert_paragraphs_at_marker(
        document,
        "[[MARKER]]",
        [ParagraphSpec("ZONA DE **CONSULTORIO**", style="List Number", highlight="yellow")],
    )
    document.save(path)

    reopened = Document(path)
    paragraph = next(paragraph for paragraph in reopened.paragraphs if "ZONA DE" in paragraph.text)
    assert any(run.text == "CONSULTORIO" and run.bold for run in paragraph.runs)
    assert any(run.font.highlight_color for run in paragraph.runs)
    with ZipFile(path) as docx_zip:
        document_xml = docx_zip.read("word/document.xml").decode("utf-8")
        numbering_xml = docx_zip.read("word/numbering.xml").decode("utf-8")
    assert "<w:numPr>" in document_xml
    assert '<w:numFmt w:val="decimal"' in numbering_xml


def test_marker_insertion_can_start_following_section_with_margin(workspace_tmp_path) -> None:
    path = workspace_tmp_path / "sectioned.docx"
    document = Document()
    document.add_paragraph("[[MARKER]]")
    insert_paragraphs_at_marker(
        document,
        "[[MARKER]]",
        [
            ParagraphSpec("Primera pagina"),
            ParagraphSpec("Segunda pagina", section_break_before=True, section_top_margin_inches=2.55),
        ],
    )
    document.save(path)

    reopened = Document(path)
    assert len(reopened.sections) == 2
    assert reopened.sections[0].top_margin.inches < 1.5
    assert reopened.sections[1].top_margin.inches >= 2.5
    assert [paragraph.text for paragraph in reopened.paragraphs if paragraph.text] == ["Primera pagina", "Segunda pagina"]


def test_marker_insertion_can_start_continuous_section(workspace_tmp_path) -> None:
    path = workspace_tmp_path / "continuous.docx"
    document = Document()
    document.add_paragraph("[[MARKER]]")
    insert_paragraphs_at_marker(
        document,
        "[[MARKER]]",
        [
            ParagraphSpec("Antes"),
            ParagraphSpec(
                "Despues",
                section_break_before=True,
                section_break_type="continuous",
                section_top_margin_inches=2.15,
            ),
        ],
    )
    document.save(path)

    reopened = Document(path)
    assert len(reopened.sections) == 2
    assert reopened.sections[0].top_margin.inches < 1.5
    assert 2.0 <= reopened.sections[1].top_margin.inches <= 2.2
    with ZipFile(path) as docx_zip:
        document_xml = docx_zip.read("word/document.xml").decode("utf-8")
    assert 'w:type w:val="continuous"' in document_xml


def test_replace_placeholders_preserves_run_formatting() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Centro: ")
    bold_run = paragraph.add_run("{{CENTER_NAME}}")
    bold_run.bold = True

    replace_placeholders(document, {"CENTER_NAME": "Hospital"})

    assert paragraph.text == "Centro: Hospital"
    assert paragraph.runs[1].text == "Hospital"
    assert paragraph.runs[1].bold is True


def test_replace_placeholders_handles_remaining_split_run_placeholder() -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("{{LOCATION}}")
    paragraph.add_run(", CHIAPAS A {{")
    paragraph.add_run("DATE}}")

    replace_placeholders(document, {"LOCATION": "TZIMOL", "DATE": "27 DE NOVIEMBRE DE 2025"})

    assert paragraph.text == "TZIMOL, CHIAPAS A 27 DE NOVIEMBRE DE 2025"


def test_marker_insertion_raises_for_missing_marker() -> None:
    document = Document()
    document.add_paragraph("Sin marcador")
    with pytest.raises(MarkerNotFoundError):
        insert_paragraphs_at_marker(document, "[[MISSING]]", [ParagraphSpec("Nada")])
