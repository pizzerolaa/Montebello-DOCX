from __future__ import annotations

from base64 import b64decode
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from lxml import etree

from app.config import Settings
from app.models.equipment import Equipment
from app.models.project import Project
from app.models.signature import Signature
from app.models.source_document import SourceDocument
from app.models.work_item import EquipmentWorkItem
from app.services.acta_generator import generate_acta


def test_generate_acta_replaces_markers_and_reopens(workspace_tmp_path) -> None:
    project = Project(
        id="project-id",
        name="Proyecto ACTA",
        location="Tapachula",
        service_date_raw="10 DE JUNIO DE 2026",
        order_number="ABC-123",
        center_name="Hospital General",
    )
    project.signatures.append(Signature(project_id=project.id, role="DELIVERER", name="Entrega", position="Cargo"))
    equipment = Equipment(
        project_id=project.id,
        sequence=1,
        equipment_type="PACKAGE",
        zone="EQUIPO 01",
        brand="YORK",
        capacity="20 TON",
        serial="S/S",
    )
    equipment.work_items.append(EquipmentWorkItem(sequence=1, title="BANDA", description="SUSTITUIDA."))
    project.equipment.append(equipment)
    project.equipment.append(
        Equipment(
            project_id=project.id,
            sequence=2,
            equipment_type="MINISPLIT",
            zone="CONSULTORIO",
            brand="YORK",
            capacity="1 TON",
            serial="MS001",
        )
    )

    output = generate_acta(project, Settings(storage_root=workspace_tmp_path / "storage"), workspace_tmp_path / "acta.docx")
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "[[SERVICE_SUMMARY]]" not in text
    assert "[[EQUIPMENT_SECTIONS]]" not in text
    assert "ABC-123" in text
    assert "1.  ZONA DE EQUIPO 01" in text
    assert "2.  ZONA DE CONSULTORIO" in text
    assert "Contenido editable" not in text
    assert "{COUNT_02}" not in text
    assert "MANTENIMIENTO CORRECTIVO DE 01 EQUIPO DE AIRE ACONDICIONADO, PAQUETE" in text
    assert "SE REALIZÓ UNA INSPECCIÓN VISUAL DEL EQUIPO DE AIRE ACONDICIONADO TIPO PAQUETE" in text
    assert not any(paragraph.text.startswith("- ") for paragraph in document.paragraphs)
    assert "MARCA: YORK" in text
    assert "CAPACIDAD: 20 TON" in text
    assert "NUMERO DE SERIE: S/N" in text
    assert len(document.sections) == 2
    assert document.sections[0].top_margin.inches < 1.5
    assert 1.2 <= document.sections[0].bottom_margin.inches <= 1.3
    assert 2.0 <= document.sections[1].top_margin.inches <= 2.2
    assert 1.2 <= document.sections[1].bottom_margin.inches <= 1.3
    assert all(run.font.name == "Arial" for paragraph in document.paragraphs for run in paragraph.runs if run.text)
    assert any(
        run.text == "Hospital General" and run.bold
        for paragraph in document.paragraphs
        for run in paragraph.runs
    )
    with ZipFile(output) as docx_zip:
        document_xml = docx_zip.read("word/document.xml").decode("utf-8")
        numbering_xml = docx_zip.read("word/numbering.xml").decode("utf-8")
    assert 'w:ascii="Arial"' in document_xml
    assert 'w:cs="Arial"' in document_xml
    assert 'w:type w:val="continuous"' in document_xml
    assert 'w:type="page"' not in document_xml
    assert "<w:numPr>" in document_xml
    assert '<w:highlight w:val="yellow"' in document_xml
    assert '<w:numFmt w:val="bullet"' in numbering_xml
    assert any(run.text.endswith(":") and run.bold for paragraph in document.paragraphs for run in paragraph.runs)
    equipment_heading = next(paragraph for paragraph in document.paragraphs if paragraph.text == "1.  ZONA DE EQUIPO 01")
    assert not any(run.font.highlight_color for run in equipment_heading.runs)
    assert not any(equipment_heading._p.findall(".//w:numPr", equipment_heading._p.nsmap))


def test_generate_acta_appends_source_report_with_tables_and_images(workspace_tmp_path) -> None:
    source_path = workspace_tmp_path / "source_report.docx"
    image_path = workspace_tmp_path / "pixel.png"
    image_path.write_bytes(
        b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    source_document = Document()
    source_document.add_paragraph("ACTA SOURCE PREAMBLE")
    source_document.add_paragraph("REPORTE DE TRABAJO")
    source_document.add_paragraph("CONTENIDO DEL REPORTE CONSERVADO")
    table = source_document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.cell(0, 0).text = "TABLA CONSERVADA"
    no_spacing = source_document.styles.add_style("Sinespaciado", WD_STYLE_TYPE.PARAGRAPH)
    no_spacing.paragraph_format.space_after = Pt(18)
    table.cell(0, 0).paragraphs[0].style = no_spacing
    table.cell(0, 0).add_paragraph("ACTIVIDADES REALIZADAS", style=no_spacing)
    signature_table = source_document.add_table(rows=1, cols=2)
    signature_table.style = "Table Grid"
    signature_table.cell(0, 0).text = "Entrega\n\n________________________________\nENTREGA"
    signature_table.cell(0, 1).text = "Recibió\n\n________________________________\nRECIBE"
    signature_table._tbl.tblPr.append(etree.fromstring(
        '<w:tblpPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'w:vertAnchor="text" w:horzAnchor="margin" w:tblpY="454"/>'
    ))
    source_document.add_paragraph()
    source_document.add_paragraph()
    source_document.add_paragraph("REPORTE FOTOGRÁFICO")
    source_document.add_picture(str(image_path))
    source_document.save(source_path)

    project = Project(
        id="project-id",
        name="Proyecto con reporte",
        location="Tuxtla Gutierrez",
        service_date_raw="02 DE DICIEMBRE DE 2025",
        order_number="OTRF-010-25",
        center_name="Clinica de Prueba",
    )
    project.source_documents.append(
        SourceDocument(
            project_id=project.id,
            original_filename="source_report.docx",
            safe_filename="source_report.docx",
            stored_path=str(source_path),
        )
    )
    project.equipment.append(
        Equipment(
            project_id=project.id,
            sequence=1,
            equipment_type="MINISPLIT",
            zone="CUARTO DE TRABAJO",
            brand="CARRIER",
            capacity="1 TON",
            serial="340C436120511210170044",
        )
    )

    output = generate_acta(project, Settings(storage_root=workspace_tmp_path / "storage"), workspace_tmp_path / "acta.docx")
    document = Document(output)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for doc_table in document.tables for row in doc_table.rows for cell in row.cells)

    assert "ACTA SOURCE PREAMBLE" not in paragraph_text
    assert "REPORTE DE TRABAJO" in paragraph_text
    assert "CONTENIDO DEL REPORTE CONSERVADO" in paragraph_text
    assert "TABLA CONSERVADA" in table_text
    assert len(document.sections) == 4
    assert 2.0 <= document.sections[1].top_margin.inches <= 2.2
    assert document.sections[2].top_margin.inches < 1.5
    assert 1.2 <= document.sections[3].top_margin.inches <= 1.3
    with ZipFile(output) as docx_zip:
        document_xml_bytes = docx_zip.read("word/document.xml")
        document_xml = document_xml_bytes.decode("utf-8")
        names = set(docx_zip.namelist())
    assert 'w:type w:val="nextPage"' in document_xml
    assert any(name.startswith("word/media/appended_report_") for name in names)

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = etree.fromstring(document_xml_bytes)
    activities_table = next(
        table
        for table in root.xpath(".//w:tbl", namespaces=namespace)
        if "ACTIVIDADES REALIZADAS" in " ".join(table.xpath(".//w:t/text()", namespaces=namespace))
    )
    assert activities_table.find("w:tblPr/w:tblStyle", namespace) is None
    assert len(activities_table.xpath("./w:tblPr/w:tblBorders/*", namespaces=namespace)) == 6
    for paragraph in activities_table.xpath(".//w:p", namespaces=namespace):
        assert paragraph.find("w:pPr/w:pStyle", namespace) is None
        spacing = paragraph.find("w:pPr/w:spacing", namespace)
        assert spacing is not None
        assert spacing.get(f"{{{namespace['w']}}}before") == "0"
        assert spacing.get(f"{{{namespace['w']}}}after") == "0"
        assert spacing.get(f"{{{namespace['w']}}}line") == "240"

    signature_table = next(
        table
        for table in root.xpath(".//w:tbl", namespaces=namespace)
        if "Entrega" in " ".join(table.xpath(".//w:t/text()", namespaces=namespace))
    )
    assert signature_table.find("w:tblPr/w:tblStyle", namespace) is None
    assert signature_table.find("w:tblPr/w:tblpPr", namespace) is None
    assert signature_table.find("w:tblPr/w:tblW", namespace).get(f"{{{namespace['w']}}}type") == "dxa"
    assert signature_table.find("w:tr/w:trPr/w:cantSplit", namespace) is not None

    body_children = list(root.find("w:body", namespace))
    signature_index = body_children.index(signature_table)
    photo_index = next(
        index
        for index, element in enumerate(body_children)
        if "REPORTE FOTOGRÁFICO" in " ".join(element.xpath(".//w:t/text()", namespaces=namespace))
    )
    assert photo_index == signature_index + 2
