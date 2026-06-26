from __future__ import annotations

from docx import Document

from app.services.docx_reader import canonical_text, read_docx_blocks


def test_read_docx_blocks_preserves_paragraph_table_header_footer_order(workspace_tmp_path) -> None:
    path = workspace_tmp_path / "ordered.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "ENCABEZADO"
    document.add_paragraph("PRIMERO")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "CELDA A"
    table.cell(0, 1).text = "CELDA B"
    document.add_paragraph("ULTIMO")
    document.sections[0].footer.paragraphs[0].text = "PIE"
    document.save(path)

    blocks = read_docx_blocks(path)
    assert [(block.kind, block.text) for block in blocks] == [
        ("paragraph", "PRIMERO"),
        ("table_cell", "CELDA A"),
        ("table_cell", "CELDA B"),
        ("paragraph", "ULTIMO"),
        ("header", "ENCABEZADO"),
        ("footer", "PIE"),
    ]


def test_canonical_text_is_accent_insensitive() -> None:
    assert canonical_text("Recepción técnico") == "RECEPCION TECNICO"
