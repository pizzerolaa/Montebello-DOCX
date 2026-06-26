from __future__ import annotations

from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION_START

from app.config import Settings
from app.models.equipment import Equipment
from app.models.project import Project
from app.services.annex_generator import generate_annex


def test_generate_annex_replaces_marker_and_reopens(workspace_tmp_path) -> None:
    project = Project(
        id="project-id",
        name="Proyecto ANEXO",
        center_name="Hospital",
        order_number="ABC-123",
        contract_date_raw="19 DE NOVIEMBRE DE 2025",
    )
    project.equipment.append(
        Equipment(
            id="eq-minisplit",
            project_id=project.id,
            sequence=1,
            equipment_type="MINISPLIT",
            zone="CONSULTORIO",
            brand="YORK",
            capacity="1 TON",
            serial="MS001",
        )
    )
    project.equipment.append(
        Equipment(
            id="eq-package",
            project_id=project.id,
            sequence=2,
            equipment_type="PACKAGE",
            zone="AZOTEA",
            brand="CARRIER",
            capacity="20 TON",
            serial="S/S",
        )
    )

    output = generate_annex(project, Settings(storage_root=workspace_tmp_path / "storage"), workspace_tmp_path / "annex.docx")
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "[[ANNEX_SECTIONS]]" not in text
    assert "ANEXO DEL REPORTE DE TRABAJO  " in text
    assert "MS001" in text
    assert "S/N" in text
    assert "Contenido editable" not in text
    assert len(document.sections) == 4
    assert document.sections[0].top_margin.inches < 1.5
    assert 1.2 <= document.sections[0].bottom_margin.inches <= 1.3
    assert 2.0 <= document.sections[1].top_margin.inches <= 2.2
    assert 1.2 <= document.sections[1].bottom_margin.inches <= 1.3
    assert document.sections[2].top_margin.inches < 1.5
    assert 1.2 <= document.sections[2].bottom_margin.inches <= 1.3
    assert 2.0 <= document.sections[3].top_margin.inches <= 2.2
    assert 1.2 <= document.sections[3].bottom_margin.inches <= 1.3
    assert document.sections[0].start_type == WD_SECTION_START.NEW_PAGE
    assert document.sections[1].start_type == WD_SECTION_START.NEW_PAGE
    assert document.sections[2].start_type == WD_SECTION_START.NEW_PAGE
    assert "AL PRINCIPIO SE REALIZÓ UNA INSPECCIÓN VISUAL DEL EQUIPO TIPO MINISPLIT" in text
    assert "EN EL MARCO DEL CONTRATO DE PRESTACIÓN DE SERVICIOS ABC-123" in text
    assert "DERIVADO DE LA EVALUACIÓN Y DIAGNÓSTICO TÉCNICO" in text
    assert "PIEZAS SUSTITUIDAS Y CORRECCIONES REALIZADAS" in text
    assert "ATENTAMENTE" in text
    assert "PEDIDO:" not in text
    assert any(run.text == "ABC-123" and run.bold for paragraph in document.paragraphs for run in paragraph.runs)
    assert any(run.text.endswith(":") and run.bold for paragraph in document.paragraphs for run in paragraph.runs)
    with ZipFile(output) as docx_zip:
        document_xml = docx_zip.read("word/document.xml").decode("utf-8")
        numbering_xml = docx_zip.read("word/numbering.xml").decode("utf-8")
    assert "w:sectPr" in document_xml
    assert 'w:type w:val="continuous"' not in document_xml
    assert '<w:numFmt w:val="bullet"' in numbering_xml
