from __future__ import annotations

from docx import Document

from app.config import Settings
from app.models.equipment import Equipment
from app.models.project import Project
from app.services.acta_generator import generate_acta
from app.services.annex_generator import generate_annex


def test_generate_acta_includes_cold_room_sections(workspace_tmp_path) -> None:
    project = Project(
        id="cold-room-project",
        name="Proyecto camara fria",
        location="Tuxtla Gutierrez",
        service_date_raw="02 DE DICIEMBRE DE 2025",
        order_number="OTRF-010-25",
        center_name="Almacen de red de frio",
    )
    project.equipment.append(
        Equipment(
            id="cold-room-1",
            project_id=project.id,
            sequence=1,
            equipment_type="COLD_ROOM",
            zone="CAMARA DE REFRIGERACION (PRECAMARA)",
            brand="BOHN",
            serial="M21K00323",
        )
    )

    output = generate_acta(project, Settings(storage_root=workspace_tmp_path / "storage"), workspace_tmp_path / "acta.docx")
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "01 CAMARA FRIA" in text
    assert "ZONA DE CAMARA DE REFRIGERACION (PRECAMARA)" in text
    assert "EQUIPO TIPO CAMARA FRIA" in text
    assert "PIEZAS SUSTITUIDAS Y CORRECCIONES REALIZADAS" in text


def test_generate_annex_includes_cold_room_piece_options(workspace_tmp_path) -> None:
    project = Project(
        id="cold-room-project",
        name="Proyecto ANEXO camara fria",
        center_name="Almacen de red de frio",
        order_number="OTRF-010-25",
        contract_date_raw="19 DE NOVIEMBRE DE 2025",
    )
    project.equipment.append(
        Equipment(
            id="cold-room-1",
            project_id=project.id,
            sequence=1,
            equipment_type="COLD_ROOM",
            zone="CAMARA DE REFRIGERACION 1",
            brand="BOHN",
            serial="M21K01361",
        )
    )

    output = generate_annex(project, Settings(storage_root=workspace_tmp_path / "storage"), workspace_tmp_path / "annex.docx")
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "TIPO CAMARA FRIA" in text
    assert "CAMARA DE REFRIGERACION 1" in text
    assert "PIEZAS SUSTITUIDAS Y CORRECCIONES REALIZADAS" in text
    assert len([paragraph for paragraph in document.paragraphs if ":" in paragraph.text]) >= 3
