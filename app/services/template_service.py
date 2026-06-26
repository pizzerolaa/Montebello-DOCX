from __future__ import annotations

from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.exceptions import MarkerNotFoundError


@dataclass(frozen=True)
class ParagraphSpec:
    text: str
    bold: bool = False
    underline: bool = False
    alignment: str = "left"
    page_break_before: bool = False
    section_break_before: bool = False
    section_break_type: str = "next_page"
    section_top_margin_inches: float | None = None
    style: str | None = None
    number_sequence: str | None = None
    font_size_pt: float = 10
    space_before_pt: float = 0
    space_after_pt: float = 0
    line_spacing: float = 1.0
    keep_with_next: bool = False
    highlight: str | None = None


def load_template(path: Path) -> DocxDocument:
    if not path.exists():
        raise FileNotFoundError(f"No se encontro la plantilla: {path}")
    return Document(path)


def replace_placeholders(document: DocxDocument, values: dict[str, str]) -> None:
    """Replace simple {{PLACEHOLDER}} values without losing run formatting."""
    for paragraph in iter_paragraphs(document, include_headers=True):
        _replace_placeholders_in_paragraph(paragraph, values)


def insert_paragraphs_at_marker(document: DocxDocument, marker: str, paragraphs: list[ParagraphSpec]) -> None:
    """Insert body paragraphs before a marker paragraph, then remove the marker."""
    marker_paragraph = find_marker_paragraph(document, marker)
    if marker_paragraph is None:
        raise MarkerNotFoundError(f"No se encontro el marcador requerido {marker}.")
    numbering_sequences: dict[str, int] = {}
    for spec in paragraphs:
        if spec.section_break_before:
            _insert_section_break_before(
                document,
                marker_paragraph,
                spec.section_top_margin_inches,
                spec.section_break_type,
            )
        inserted = marker_paragraph.insert_paragraph_before()
        if spec.style:
            _apply_style_or_numbering(inserted, spec.style, spec.number_sequence, numbering_sequences)
        if spec.page_break_before:
            inserted.add_run().add_break(WD_BREAK.PAGE)
        _add_marked_runs(inserted, spec)
        inserted.alignment = _alignment(spec.alignment)
        paragraph_format = inserted.paragraph_format
        paragraph_format.space_before = Pt(spec.space_before_pt)
        paragraph_format.space_after = Pt(spec.space_after_pt)
        paragraph_format.line_spacing = spec.line_spacing
        if spec.bold or spec.keep_with_next:
            paragraph_format.keep_with_next = True
    _remove_paragraph(marker_paragraph)


def _insert_section_break_before(
    document: DocxDocument,
    marker_paragraph: Paragraph,
    section_top_margin_inches: float | None,
    section_break_type: str,
) -> None:
    """Start a new section before the marker while preserving the current section layout."""
    break_paragraph = marker_paragraph.insert_paragraph_before()
    ppr = break_paragraph._p.get_or_add_pPr()
    section_properties = deepcopy(document.sections[-1]._sectPr)
    if section_break_type == "continuous":
        section_type = section_properties.find(qn("w:type"))
        if section_type is None:
            section_type = OxmlElement("w:type")
            section_properties.insert(0, section_type)
        section_type.set(qn("w:val"), "continuous")
    ppr.append(section_properties)
    if section_top_margin_inches is not None:
        document.sections[-1].top_margin = Inches(section_top_margin_inches)


def normalize_document_font(document: DocxDocument, font_name: str = "Arial") -> None:
    """Force template and generated text to use the configured Word font."""
    for style in document.styles:
        if hasattr(style, "font"):
            style.font.name = font_name
            if style.font.size is None:
                style.font.size = Pt(10)
            if style.element.rPr is not None and style.element.rPr.rFonts is not None:
                _set_rfonts(style.element.rPr.rFonts, font_name)
    for paragraph in iter_paragraphs(document, include_headers=True):
        for run in paragraph.runs:
            set_run_font(run, font_name=font_name)


def find_marker_paragraph(document: DocxDocument, marker: str) -> Paragraph | None:
    for paragraph in iter_paragraphs(document):
        if marker in paragraph.text:
            return paragraph
    return None


def iter_paragraphs(document: DocxDocument, *, include_headers: bool = False):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    if include_headers:
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph
            for table in section.header.tables:
                yield from _iter_table_paragraphs(table)
            for paragraph in section.footer.paragraphs:
                yield paragraph
            for table in section.footer.tables:
                yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _alignment(value: str):
    if value == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    if value == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    return WD_ALIGN_PARAGRAPH.LEFT


def _apply_style_or_numbering(
    paragraph: Paragraph,
    style: str,
    number_sequence: str | None,
    numbering_sequences: dict[str, int],
) -> None:
    if style == "List Number":
        _apply_numbering(paragraph, "decimal", number_sequence, numbering_sequences)
        return
    if style == "List Bullet":
        _apply_numbering(paragraph, "bullet", number_sequence, numbering_sequences)
        return
    try:
        paragraph.style = style
    except KeyError:
        return


def _apply_numbering(
    paragraph: Paragraph,
    num_format: str,
    number_sequence: str | None,
    numbering_sequences: dict[str, int],
) -> None:
    sequence_key = f"{num_format}:{number_sequence or 'default'}"
    num_id = numbering_sequences.get(sequence_key)
    if num_id is None:
        num_id = _create_numbering(paragraph.part.numbering_part._element, num_format)
        numbering_sequences[sequence_key] = num_id
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), str(num_id))


def _create_numbering(numbering, num_format: str) -> int:
    abstract_id = _next_numbering_id(numbering, "abstractNum", "abstractNumId")
    num_id = _next_numbering_id(numbering, "num", "numId")

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level_type = OxmlElement("w:multiLevelType")
    multi_level_type.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi_level_type)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_format)
    lvl.append(fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if num_format == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    lvl.append(ppr)
    abstract_num.append(lvl)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def _next_numbering_id(numbering, element_name: str, attribute_name: str) -> int:
    ids = [
        int(element.get(qn(f"w:{attribute_name}")))
        for element in numbering.findall(qn(f"w:{element_name}"))
        if element.get(qn(f"w:{attribute_name}")) is not None
    ]
    return max(ids, default=0) + 1


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def _replace_placeholders_in_paragraph(paragraph: Paragraph, values: dict[str, str]) -> None:
    if not paragraph.runs:
        return
    replaced_in_place = False
    for run in paragraph.runs:
        updated = run.text
        for placeholder, value in values.items():
            updated = updated.replace(f"{{{{{placeholder}}}}}", value or "")
        if updated != run.text:
            run.text = updated
            set_run_font(run)
            replaced_in_place = True
    if replaced_in_place and not _has_placeholders(paragraph.text, values):
        return

    original = "".join(run.text for run in paragraph.runs)
    updated = original
    for placeholder, value in values.items():
        updated = updated.replace(f"{{{{{placeholder}}}}}", value or "")
    if updated == original:
        return
    paragraph.runs[0].text = updated
    set_run_font(paragraph.runs[0])
    for run in paragraph.runs[1:]:
        run.text = ""


def set_run_font(run: Run, *, font_name: str = "Arial", size_pt: float | None = None) -> None:
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    _set_rfonts(rfonts, font_name)


def _add_marked_runs(paragraph: Paragraph, spec: ParagraphSpec) -> None:
    parts = spec.text.split("**")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = spec.bold or index % 2 == 1
        run.underline = spec.underline
        if spec.highlight == "yellow":
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        set_run_font(run, size_pt=spec.font_size_pt)


def _set_rfonts(rfonts, font_name: str) -> None:
    for attribute in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attribute}"), font_name)


def _has_placeholders(text: str, values: dict[str, str]) -> bool:
    return any(f"{{{{{placeholder}}}}}" in text for placeholder in values)
