from __future__ import annotations

from dataclasses import replace
from pathlib import Path

from docx import Document

from app.config import Settings
from app.models.equipment import Equipment
from app.models.project import Project
from app.services.acta_generator import _apply_document_layout, _content_specs, _piece_spec
from app.services.equipment_format import normalize_serial
from app.services.excel_service import derived_work_done
from app.services.template_service import ParagraphSpec, insert_paragraphs_at_marker, load_template, normalize_document_font
from app.services.work_option_service import cold_room_work_for_equipment, minisplit_work_for_equipment, package_work_for_equipment


FIRST_PAGE_TOP_MARGIN_INCHES = 0.98
CONTINUATION_TOP_MARGIN_INCHES = 2.15
CONTINUATION_BREAK_INDEX = {
    "MINISPLIT": 10,
    "PACKAGE": 9,
    "COLD_ROOM": 9,
}


def generate_annex(project: Project, settings: Settings, output_path: Path) -> Path:
    """Generate ANEXO DOCX from the configured Word template marker."""
    template_path = Path(__file__).resolve().parents[1] / "word_templates" / "annex_template.docx"
    document = load_template(template_path)
    _apply_document_layout(document)
    if document.paragraphs and not document.paragraphs[0].text.strip():
        paragraph = document.paragraphs[0]._element
        paragraph.getparent().remove(paragraph)
    insert_paragraphs_at_marker(document, "[[ANNEX_SECTIONS]]", build_annex_sections(project))
    normalize_document_font(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    Document(output_path)
    return output_path


def build_annex_sections(project: Project) -> list[ParagraphSpec]:
    paragraphs: list[ParagraphSpec] = []
    for index, equipment in enumerate(sorted(project.equipment, key=lambda item: item.sequence)):
        builder = ANNEX_BUILDERS.get(equipment.equipment_type, build_unknown_annex_section)
        section = builder(project, equipment)
        if index > 0:
            section[0] = replace(
                section[0],
                section_break_before=True,
                section_top_margin_inches=FIRST_PAGE_TOP_MARGIN_INCHES,
            )
        break_index = min(
            CONTINUATION_BREAK_INDEX.get(equipment.equipment_type, max(len(section) - 5, 1)),
            len(section) - 1,
        )
        section[break_index] = replace(
            section[break_index],
            section_break_before=True,
            section_top_margin_inches=CONTINUATION_TOP_MARGIN_INCHES,
        )
        paragraphs.extend(section)
    return paragraphs


def build_minisplit_annex_section(project: Project, equipment: Equipment) -> list[ParagraphSpec]:
    paragraphs = _annex_header(project, equipment)
    paragraphs.extend(_annex_content_specs("minisplit", "annex_maintenance.txt"))
    paragraphs.extend(_work_specs(equipment))
    paragraphs.extend(_signature_specs())
    return paragraphs


def build_package_annex_section(project: Project, equipment: Equipment) -> list[ParagraphSpec]:
    paragraphs = _annex_header(project, equipment)
    paragraphs.extend(_annex_content_specs("package", "annex_maintenance.txt"))
    paragraphs.extend(_work_specs(equipment))
    paragraphs.extend(_signature_specs())
    return paragraphs


def build_cold_room_annex_section(project: Project, equipment: Equipment) -> list[ParagraphSpec]:
    paragraphs = _annex_header(project, equipment)
    paragraphs.extend(_annex_content_specs("cold_room", "annex_maintenance.txt"))
    paragraphs.extend(_work_specs(equipment))
    paragraphs.extend(_signature_specs())
    return paragraphs


def build_unknown_annex_section(project: Project, equipment: Equipment) -> list[ParagraphSpec]:
    paragraphs = _annex_header(project, equipment)
    paragraphs.extend(_work_specs(equipment))
    paragraphs.extend(_signature_specs())
    return paragraphs


ANNEX_BUILDERS = {
    "MINISPLIT": build_minisplit_annex_section,
    "PACKAGE": build_package_annex_section,
    "COLD_ROOM": build_cold_room_annex_section,
}


def _annex_header(project: Project, equipment: Equipment) -> list[ParagraphSpec]:
    order_number = project.order_number or ""
    contract_date = project.contract_date_raw or project.service_date_raw or ""
    center_name = project.center_name or ""
    equipment_type = _equipment_type_label(equipment)
    brand = equipment.brand or ""
    serial = normalize_serial(equipment.serial)
    capacity = f", DE CAPACIDAD **{equipment.capacity}**" if equipment.capacity else ""
    zone = equipment.zone or ""
    return [
        ParagraphSpec(
            "ANEXO DEL REPORTE DE TRABAJO  ",
            bold=True,
            underline=True,
            alignment="center",
            font_size_pt=12,
            space_after_pt=54,
        ),
        ParagraphSpec(center_name, bold=True, alignment="center", font_size_pt=11),
        ParagraphSpec(
            f"DETALLE TECNICO DE LOS TRABAJOS REALIZADOS PARA EL MANTENIMIENTO CORRECTIVO DERIVADOS DEL CONTRATO {order_number}",
            bold=True,
            alignment="center",
            font_size_pt=11,
            space_after_pt=14,
        ),
        ParagraphSpec(
            (
                "EN EL MARCO DEL CONTRATO DE PRESTACIÓN DE SERVICIOS "
                f"**{order_number}**, DE FECHA {contract_date}, CELEBRADO ENTRE EL H. "
                "INSTITUTO DE SALUD Y MI REPRESENTADA, **MONTEBELLO ESTRUCTURAL S.A. DE C.V.**, "
                "DICHO INSTRUMENTO AMPARA LA EJECUCIÓN DE LOS TRABAJOS RELATIVOS A LA "
                '**PARTIDA 35701.- MANTENIMIENTO Y CONSERVACIÓN DE MAQUINARIA Y EQUIPO**, ANEXO 1.5 (REQUISICIÓN 17).'
            ),
            alignment="justify",
            space_after_pt=10,
        ),
        ParagraphSpec(
            (
                "DERIVADO DE LA EVALUACIÓN Y DIAGNÓSTICO TÉCNICO REALIZADO A LA INFRAESTRUCTURA DEL "
                f"**{center_name}**, SE ENCONTRÓ EL EQUIPO DE AIRE ACONDICIONADO {equipment_type} "
                f"MARCA **{brand}**, NÚMERO DE SERIE **{serial}**{capacity}, "
                f"UBICADO EN LA ZONA DE **{zone}**."
            ),
            alignment="justify",
            space_after_pt=10,
        ),
    ]


def _work_specs(equipment: Equipment) -> list[ParagraphSpec]:
    if equipment.equipment_type == "MINISPLIT":
        work_done = minisplit_work_for_equipment(equipment) or derived_work_done(equipment)
        if not work_done:
            return []
        return [
            *_annex_content_specs("minisplit", "work_intro.txt"),
            ParagraphSpec(work_done, alignment="justify", space_after_pt=18),
        ]
    if equipment.equipment_type == "PACKAGE":
        work_items = package_work_for_equipment(equipment)
        if work_items:
            return [
                ParagraphSpec(
                    "PIEZAS SUSTITUIDAS Y CORRECCIONES REALIZADAS",
                    bold=True,
                    alignment="justify",
                    space_before_pt=14,
                    space_after_pt=12,
                ),
                *_annex_content_specs("package", "work_intro.txt"),
                *[_piece_spec(line, number_sequence=f"work-{equipment.sequence}") for line in work_items],
                *_annex_content_specs("package", "final_note.txt"),
            ]
    if equipment.equipment_type == "COLD_ROOM":
        work_items = cold_room_work_for_equipment(equipment)
        if work_items:
            return [
                ParagraphSpec(
                    "PIEZAS SUSTITUIDAS Y CORRECCIONES REALIZADAS",
                    bold=True,
                    alignment="justify",
                    space_before_pt=14,
                    space_after_pt=12,
                ),
                *_annex_content_specs("cold_room", "work_intro.txt"),
                *[_piece_spec(line, number_sequence=f"work-{equipment.sequence}") for line in work_items],
                *_annex_content_specs("cold_room", "final_note.txt"),
            ]
    work_done = derived_work_done(equipment)
    if not work_done:
        return []
    return [ParagraphSpec(work_done, alignment="justify", space_after_pt=18)]


def _signature_specs() -> list[ParagraphSpec]:
    return [
        ParagraphSpec("ATENTAMENTE", bold=True, alignment="center", font_size_pt=11, space_before_pt=24, space_after_pt=36),
        ParagraphSpec("________________________________", bold=True, alignment="center"),
        ParagraphSpec("C. IAN LUIS OSORNO HERNANDEZ", bold=True, alignment="center"),
        ParagraphSpec("ADMINISTRADOR UNICO", bold=True, alignment="center"),
        ParagraphSpec("MONTEBELLO ESTRUCTURAL S.A. DE C.V.", bold=True, alignment="center", space_after_pt=8),
    ]


def _annex_content_specs(equipment_folder: str, filename: str) -> list[ParagraphSpec]:
    return [
        replace(
            spec,
            font_size_pt=min(spec.font_size_pt, 9.5),
            space_after_pt=min(spec.space_after_pt, 6),
        )
        for spec in _content_specs(equipment_folder, filename)
    ]


def _equipment_type_label(equipment: Equipment) -> str:
    if equipment.equipment_type == "MINISPLIT":
        return "TIPO MINISPLIT"
    if equipment.equipment_type == "PACKAGE":
        return "TIPO PAQUETE"
    if equipment.equipment_type == "COLD_ROOM":
        return "TIPO CAMARA FRIA"
    return equipment.equipment_type or "TIPO PENDIENTE"
