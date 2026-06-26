from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


@dataclass(frozen=True)
class DocumentBlock:
    kind: Literal["paragraph", "table_cell", "header", "footer"]
    text: str
    order: int
    source_name: str


def normalize_spaces(value: str) -> str:
    """Collapse repeated whitespace without changing meaning."""
    return re.sub(r"\s+", " ", value).strip()


def remove_accents(value: str) -> str:
    """Remove diacritic marks for accent-insensitive comparisons."""
    normalized = unicodedata.normalize("NFD", value)
    return "".join(char for char in normalized if unicodedata.category(char) != "Mn")


def canonical_text(value: str) -> str:
    """Normalize text for comparison while preserving original storage values elsewhere."""
    cleaned = normalize_spaces(remove_accents(value))
    cleaned = cleaned.replace("：", ":").replace("–", "-").replace("—", "-")
    return cleaned.upper()


def clean_label_value(value: str) -> str:
    """Clean a label value extracted from Word text."""
    return normalize_spaces(value.strip(" :\t\r\n"))


def _iter_body_items(parent: DocxDocument | _Cell):
    body = parent.element.body if isinstance(parent, DocxDocument) else parent._tc
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _append_text_block(
    blocks: list[DocumentBlock],
    kind: Literal["paragraph", "table_cell", "header", "footer"],
    text: str,
    source_name: str,
) -> None:
    if text and normalize_spaces(text):
        blocks.append(DocumentBlock(kind=kind, text=text, order=len(blocks), source_name=source_name))


def read_docx_blocks(path: Path | str, source_name: str | None = None) -> list[DocumentBlock]:
    """Read DOCX text from paragraphs, tables, headers, and footers in stable document order."""
    doc_path = Path(path)
    document = Document(doc_path)
    source = source_name or doc_path.name
    blocks: list[DocumentBlock] = []

    for item in _iter_body_items(document):
        if isinstance(item, Paragraph):
            _append_text_block(blocks, "paragraph", item.text, source)
        elif isinstance(item, Table):
            for row in item.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _append_text_block(blocks, "table_cell", paragraph.text, source)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            _append_text_block(blocks, "header", paragraph.text, source)
        for paragraph in section.footer.paragraphs:
            _append_text_block(blocks, "footer", paragraph.text, source)

    return blocks
