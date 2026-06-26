from __future__ import annotations

from pathlib import Path

from docx import Document


for path in [
    Path(r"C:\Users\ferna\Documents\LOTE 3\LOTE 3\2. TZIMOL\2. TZIMOL.docx"),
    Path(r"C:\Users\ferna\Documents\LOTE 3\LOTE 3\4. NUEVA CONCORDIA PAQUETE.docx"),
]:
    document = Document(path)
    print("---", path)
    for index, paragraph in enumerate(document.paragraphs):
        text = " ".join(paragraph.text.split())
        key = text.upper()
        if any(
            marker in key
            for marker in [
                "SIN OTRO PARTICULAR",
                "REPORTE DE TRABAJO",
                "ACTA DE ENTREGA",
                "DESGLOSE DE LOS MANTENIMIENTOS",
                "REPORTE FOTOGRAFICO",
            ]
        ):
            print(index, repr(text[:220]))
    print("paragraphs", len(document.paragraphs), "tables", len(document.tables), "sections", len(document.sections))
